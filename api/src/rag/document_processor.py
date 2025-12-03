"""
Document Processor for Evidence Text Extraction

Extracts text from various document formats:
- PDF (using PyMuPDF)
- DOCX (using python-docx)
- TXT (plain text)
- Images with text (OCR via Azure Document Intelligence - optional)

This enables semantic search across evidence file contents.
"""

import os
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Extract text content from various document formats
    
    Supports:
    - PDF files (.pdf)
    - Word documents (.docx, .doc)
    - Plain text (.txt, .csv, .log)
    - Images (.png, .jpg, .jpeg) - OCR optional
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'text': ['.txt', '.csv', '.log', '.md', '.json', '.xml'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
    }
    
    def __init__(self):
        """Initialize document processor with available extractors"""
        self.extractors = {}
        self._init_extractors()
    
    def _init_extractors(self):
        """Initialize available text extractors based on installed packages"""
        
        # Try to import PyMuPDF for PDF extraction
        try:
            import fitz  # PyMuPDF
            self.extractors['pdf'] = self._extract_pdf
            logger.info("✅ PDF extractor initialized (PyMuPDF)")
        except ImportError:
            logger.warning("⚠️ PyMuPDF not installed. PDF extraction disabled. Install: pip install PyMuPDF")
            self.extractors['pdf'] = self._extract_pdf_fallback
        
        # Try to import python-docx for Word documents
        try:
            import docx
            self.extractors['word'] = self._extract_docx
            logger.info("✅ Word extractor initialized (python-docx)")
        except ImportError:
            logger.warning("⚠️ python-docx not installed. DOCX extraction disabled. Install: pip install python-docx")
            self.extractors['word'] = self._extract_text_fallback
        
        # Plain text always available
        self.extractors['text'] = self._extract_text
        logger.info("✅ Text extractor initialized")
        
        # Image OCR (optional - requires Azure Document Intelligence)
        self.extractors['image'] = self._extract_image_ocr
    
    def get_file_type(self, file_path: str) -> Optional[str]:
        """Determine file type from extension"""
        ext = Path(file_path).suffix.lower()
        
        for file_type, extensions in self.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                return file_type
        
        return None
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        return self.get_file_type(file_path) is not None
    
    async def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict with:
            - success: bool
            - text: Extracted text content
            - metadata: File metadata (pages, word count, etc.)
            - error: Error message if failed
        """
        if not os.path.exists(file_path):
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"File not found: {file_path}"
            }
        
        file_type = self.get_file_type(file_path)
        
        if not file_type:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"Unsupported file format: {Path(file_path).suffix}"
            }
        
        extractor = self.extractors.get(file_type)
        
        if not extractor:
            return {
                "success": False,
                "text": "",
                "metadata": {},
                "error": f"No extractor available for {file_type}"
            }
        
        try:
            result = await extractor(file_path)
            
            # Add common metadata
            result["metadata"]["file_path"] = file_path
            result["metadata"]["file_name"] = os.path.basename(file_path)
            result["metadata"]["file_size"] = os.path.getsize(file_path)
            result["metadata"]["file_type"] = file_type
            
            # Calculate word count
            if result["text"]:
                result["metadata"]["word_count"] = len(result["text"].split())
                result["metadata"]["char_count"] = len(result["text"])
            
            logger.info(f"✅ Extracted {result['metadata'].get('word_count', 0)} words from {file_path}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {file_path}: {e}", exc_info=True)
            return {
                "success": False,
                "text": "",
                "metadata": {"file_path": file_path},
                "error": str(e)
            }
    
    async def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF"""
        import fitz  # PyMuPDF
        
        text_parts = []
        page_texts = []
        
        with fitz.open(file_path) as doc:
            num_pages = len(doc)
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                text_parts.append(page_text)
                page_texts.append({
                    "page": page_num + 1,
                    "text": page_text,
                    "char_count": len(page_text)
                })
        
        full_text = "\n\n".join(text_parts)
        
        return {
            "success": True,
            "text": full_text,
            "metadata": {
                "page_count": num_pages,
                "pages": page_texts
            },
            "error": None
        }
    
    async def _extract_pdf_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback PDF extraction when PyMuPDF is not available"""
        logger.warning(f"PDF extraction skipped (PyMuPDF not installed): {file_path}")
        return {
            "success": False,
            "text": "",
            "metadata": {},
            "error": "PyMuPDF not installed. Install with: pip install PyMuPDF"
        }
    
    async def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx"""
        import docx
        
        doc = docx.Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
        
        full_text = "\n\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[Tables]\n" + "\n".join(table_texts)
        
        return {
            "success": True,
            "text": full_text,
            "metadata": {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            },
            "error": None
        }
    
    async def _extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files"""
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                return {
                    "success": True,
                    "text": text,
                    "metadata": {
                        "encoding": encoding,
                        "line_count": text.count('\n') + 1
                    },
                    "error": None
                }
            except UnicodeDecodeError:
                continue
        
        return {
            "success": False,
            "text": "",
            "metadata": {},
            "error": "Could not decode file with any supported encoding"
        }
    
    async def _extract_text_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback for unsupported formats - try as plain text"""
        return await self._extract_text(file_path)
    
    async def _extract_image_ocr(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from images using OCR
        
        Currently returns placeholder - can be extended with:
        - Azure Document Intelligence
        - Tesseract OCR
        - AWS Textract
        """
        logger.info(f"Image OCR requested for {file_path} - OCR not configured")
        
        return {
            "success": False,
            "text": "",
            "metadata": {
                "ocr_available": False
            },
            "error": "OCR not configured. Image text extraction requires Azure Document Intelligence or Tesseract."
        }
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats"""
        available = {}
        
        for file_type, extensions in self.SUPPORTED_EXTENSIONS.items():
            extractor = self.extractors.get(file_type)
            is_fallback = extractor and 'fallback' in extractor.__name__
            
            available[file_type] = {
                "extensions": extensions,
                "available": extractor is not None and not is_fallback,
                "note": "Fallback mode" if is_fallback else None
            }
        
        return available


# Global instance
document_processor = DocumentProcessor()
